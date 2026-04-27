import io

from docx import Document


def parse(content: bytes, filename: str) -> dict:
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(rows)
    text = "\n".join(paragraphs)
    if tables:
        flat = "\n".join("\t".join(cell for cell in row) for table in tables for row in table)
        text = f"{text}\n{flat}".strip()
    return {
        "text": text,
        "raw": {
            "format": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "tables": tables,
        },
    }
