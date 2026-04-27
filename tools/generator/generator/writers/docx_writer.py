from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from . import _text


def write(d: dict, path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_heading(d["title"], level=1)
    title.alignment = 1

    lines = _text.render_lines(d)[1:]
    for line in lines:
        if line == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    if "items" in d and d["items"]:
        rows = _text.items_table(d["items"])
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            cells = table.rows[r_idx].cells
            for c_idx, val in enumerate(row):
                cells[c_idx].text = val

    doc.save(str(path))
